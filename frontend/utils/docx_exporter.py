import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_draft_docx(sections: dict) -> io.BytesIO:
    """
    Creates a DOCX document from the IMRaD sections.
    Injects a prominent "DRAFT" disclaimer into the header.
    """
    doc = Document()
    
    # Set document title properties
    styles = doc.styles
    title_style = styles['Title']
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(24)
    
    # Add Title
    title = doc.add_heading("Research Article Draft", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add Disclaimer block
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run("DRAFT – FOR HUMAN REVIEW ONLY\nNot for Submission")
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00) # Red color
    
    doc.add_paragraph() # spacing
    
    # Standard section ordering
    order = ["Abstract", "Introduction", "Methods", "Results", "Discussion"]
    
    for section_name in order:
        content = sections.get(section_name)
        if content:
            # Heading
            heading = doc.add_heading(section_name, level=1)
            # Content
            p = doc.add_paragraph(content)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            doc.add_paragraph() # spacing
            
    # Save to BytesIO buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer
